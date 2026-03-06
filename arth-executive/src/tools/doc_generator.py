from langchain_core.tools import tool
import os
import re
import uuid
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF


def _parse_markdown_to_docx(doc: Document, content: str):
    """
    Interpreta o conteúdo em Markdown leve e aplica formatação rica no Word.
    Suporta: # Títulos, ## Subtítulos, **negrito**, - listas, texto normal.
    """
    lines = content.split('\n')
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
        
        # --- Heading 1: # Título ---
        if stripped.startswith('# ') and not stripped.startswith('## '):
            heading_text = stripped[2:].strip()
            doc.add_heading(heading_text, level=1)
        
        # --- Heading 2: ## Subtítulo ---
        elif stripped.startswith('## '):
            heading_text = stripped[3:].strip()
            doc.add_heading(heading_text, level=2)
        
        # --- Heading 3: ### Sub-subtítulo ---
        elif stripped.startswith('### '):
            heading_text = stripped[4:].strip()
            doc.add_heading(heading_text, level=3)
        
        # --- Lista com bullet: - item ou * item ---
        elif stripped.startswith('- ') or stripped.startswith('* '):
            item_text = stripped[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            _add_rich_text(p, item_text)
        
        # --- Lista numerada: 1. item ---
        elif re.match(r'^\d+\.\s', stripped):
            item_text = re.sub(r'^\d+\.\s', '', stripped).strip()
            p = doc.add_paragraph(style='List Number')
            _add_rich_text(p, item_text)
        
        # --- Separador horizontal: --- ou ___ ---
        elif stripped in ('---', '___', '***'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run('─' * 60)
            run.font.color.rgb = RGBColor(180, 180, 180)
            run.font.size = Pt(8)
        
        # --- Texto normal (com possível **negrito** inline) ---
        else:
            p = doc.add_paragraph()
            _add_rich_text(p, stripped)


def _add_rich_text(paragraph, text: str):
    """
    Adiciona texto com suporte a **negrito** inline no parágrafo.
    """
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def _parse_markdown_to_pdf(pdf: FPDF, content: str):
    """
    Interpreta o conteúdo em Markdown leve e aplica formatação rica no PDF.
    Suporta: # Títulos, ## Subtítulos, **negrito**, - listas, texto normal.
    """
    lines = content.split('\n')
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            pdf.ln(4)
            continue
        
        safe = stripped.encode('latin-1', 'replace').decode('latin-1')
        
        # --- Heading 1 ---
        if stripped.startswith('# ') and not stripped.startswith('## '):
            heading = stripped[2:].strip().encode('latin-1', 'replace').decode('latin-1')
            pdf.ln(6)
            pdf.set_font("Helvetica", style='B', size=16)
            pdf.multi_cell(0, 10, heading)
            # Linha separadora embaixo do heading
            pdf.set_draw_color(100, 100, 100)
            pdf.line(pdf.get_x() + 10, pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
            pdf.ln(4)
        
        # --- Heading 2 ---
        elif stripped.startswith('## '):
            heading = stripped[3:].strip().encode('latin-1', 'replace').decode('latin-1')
            pdf.ln(4)
            pdf.set_font("Helvetica", style='B', size=14)
            pdf.multi_cell(0, 9, heading)
            pdf.ln(2)
        
        # --- Heading 3 ---
        elif stripped.startswith('### '):
            heading = stripped[4:].strip().encode('latin-1', 'replace').decode('latin-1')
            pdf.ln(3)
            pdf.set_font("Helvetica", style='BI', size=12)
            pdf.multi_cell(0, 8, heading)
            pdf.ln(2)
        
        # --- Bullet list ---
        elif stripped.startswith('- ') or stripped.startswith('* '):
            item = stripped[2:].strip().encode('latin-1', 'replace').decode('latin-1')
            pdf.set_font("Helvetica", size=11)
            pdf.cell(10)  # indentação
            pdf.multi_cell(0, 7, f"  -  {item}")
            pdf.ln(1)
        
        # --- Numbered list ---
        elif re.match(r'^\d+\.\s', stripped):
            pdf.set_font("Helvetica", size=11)
            pdf.cell(10)  # indentação
            pdf.multi_cell(0, 7, f"  {safe}")
            pdf.ln(1)
        
        # --- Separador ---
        elif stripped in ('---', '___', '***'):
            pdf.ln(4)
            pdf.set_draw_color(180, 180, 180)
            pdf.line(pdf.get_x() + 10, pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
            pdf.ln(6)
        
        # --- Texto normal ---
        else:
            pdf.set_font("Helvetica", size=11)
            pdf.multi_cell(0, 7, safe)
            pdf.ln(2)


@tool
def generate_docx(title: str, content: str) -> str:
    """
    Cria um documento Word (.docx) PROFISSIONAL e bem formatado.
    Use esta ferramenta quando o usuário pedir para gerar um relatório,
    contrato, resumo ou qualquer texto formatado em arquivo.
    
    IMPORTANTE: O 'content' DEVE ser escrito em Markdown leve para que
    o documento fique BONITO e ORGANIZADO. Use:
    - # para títulos de seção
    - ## para subtítulos
    - ### para sub-subtítulos
    - **texto** para negrito
    - - item para listas com bullet
    - 1. item para listas numeradas
    - --- para separadores horizontais
    
    Exemplo de content BOM:
    # Introdução
    Este documento aborda **temas importantes** sobre o assunto.
    
    ## Seção 1
    - Primeiro ponto relevante
    - Segundo ponto relevante
    
    ---
    
    ## Conclusão
    Resumo final do documento.
    """
    try:
        data_dir = os.path.join(os.getcwd(), "data", "outputs")
        os.makedirs(data_dir, exist_ok=True)
        
        filename = f"{uuid.uuid4().hex[:8]}_{title.replace(' ', '_').lower()}.docx"
        filepath = os.path.join(data_dir, filename)
        
        doc = Document()
        
        # Título principal do documento (Heading 0 = título de capa)
        doc.add_heading(title, 0)
        
        # Processa o conteúdo Markdown em formatação rica
        _parse_markdown_to_docx(doc, content)
                
        doc.save(filepath)
        return (
            f"O documento Word foi gerado com formatação profissional. "
            f"Agora escreva uma BREVE APRESENTAÇÃO do documento para o usuário (2 a 4 linhas), "
            f"destacando os tópicos principais que o documento cobre. "
            f"E no FINAL da sua resposta, cole EXATAMENTE esta tag: <SEND_FILE:{filename}>\n"
            f"NÃO crie links markdown."
        )
    except Exception as e:
        return f"Erro ao gerar documento Word: {str(e)}"

@tool
def generate_pdf(title: str, content: str) -> str:
    """
    Cria um documento PDF PROFISSIONAL e bem formatado.
    Use esta ferramenta APENAS quando o usuário pedir explicitamente por um arquivo .pdf.
    
    IMPORTANTE: O 'content' DEVE ser escrito em Markdown leve para que
    o documento fique BONITO e ORGANIZADO. Use:
    - # para títulos de seção
    - ## para subtítulos
    - ### para sub-subtítulos
    - **texto** para negrito
    - - item para listas com bullet
    - 1. item para listas numeradas
    - --- para separadores horizontais
    
    Escreva o conteúdo de forma DETALHADA, RICA e ESTRUTURADA, com várias seções.
    """
    try:
        data_dir = os.path.join(os.getcwd(), "data", "outputs")
        os.makedirs(data_dir, exist_ok=True)
        
        filename = f"{uuid.uuid4().hex[:8]}_{title.replace(' ', '_').lower()}.pdf"
        filepath = os.path.join(data_dir, filename)
        
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # Título principal (grande, centralizado, bold)
        pdf.set_font("Helvetica", style='B', size=22)
        pdf.multi_cell(0, 12, title.encode('latin-1', 'replace').decode('latin-1'), align='C')
        pdf.ln(4)
        
        # Linha separadora abaixo do título
        pdf.set_draw_color(60, 60, 60)
        pdf.set_line_width(0.5)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(8)
        
        # Processa o conteúdo Markdown em formatação rica
        _parse_markdown_to_pdf(pdf, content)
                
        pdf.output(filepath)
        
        return (
            f"O documento PDF foi gerado com formatação profissional. "
            f"Agora escreva uma BREVE APRESENTAÇÃO do documento para o usuário (2 a 4 linhas), "
            f"destacando os tópicos principais que o documento cobre. "
            f"E no FINAL da sua resposta, cole EXATAMENTE esta tag: <SEND_FILE:{filename}>\n"
            f"NÃO crie links markdown."
        )
    except Exception as e:
        return f"Erro ao gerar PDF: {str(e)}"
