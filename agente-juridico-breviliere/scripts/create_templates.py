"""
Gera os templates DOCX profissionais para o Agente Jurídico Breviliere.

Uso:
    python scripts/create_templates.py

Cria 3 templates em src/documents/templates/:
  - briefing.docx       → Resumo do caso para a equipe jurídica
  - proposta_honorarios.docx → Proposta comercial para o cliente
  - contrato_servicos.docx   → Contrato de prestação de serviços

Os templates usam placeholders Jinja2 ({{ variavel }}) compatíveis
com a biblioteca docxtpl usada em src/documents/generator.py.
"""

import os
import sys
from pathlib import Path

# Adiciona raiz ao path
sys.path.insert(0, str(Path(__file__).parent.parent))

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("❌ python-docx não instalado. Execute: pip install python-docx")
    sys.exit(1)

TEMPLATES_DIR = Path(__file__).parent.parent / "src" / "documents" / "templates"


def set_cell_shading(cell, color: str):
    """Aplica cor de fundo a uma célula da tabela."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_heading(doc, text, level=1):
    """Adiciona heading com estilo personalizado."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)  # Azul escuro
    return heading


def add_info_row(table, label, value, row_idx):
    """Preenche uma linha da tabela com label e valor."""
    row = table.rows[row_idx]
    row.cells[0].text = label
    row.cells[1].text = value
    # Estilo do label
    for paragraph in row.cells[0].paragraphs:
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)


# ============================================================
# TEMPLATE 1: BRIEFING JURÍDICO
# ============================================================
def create_briefing_template():
    """Cria o template de briefing para a equipe jurídica."""
    doc = Document()

    # Configurar margens
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ---- CABEÇALHO ----
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run("BREVILIERE ADVOCACIA")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Briefing de Atendimento — Documento Interno")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("─" * 60)

    # ---- DADOS DO ATENDIMENTO ----
    add_styled_heading(doc, "1. Dados do Atendimento", level=2)

    table = doc.add_table(rows=6, cols=2, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    data = [
        ("Data do atendimento:", "{{ date }}"),
        ("Canal:", "{{ channel }}"),
        ("Protocolo:", "{{ protocol_id }}"),
        ("Score de qualificação:", "{{ score }}/100"),
        ("Urgência:", "{{ urgency }}"),
        ("Área jurídica:", "{{ area }}"),
    ]
    for i, (label, value) in enumerate(data):
        add_info_row(table, label, value, i)

    # ---- DADOS DO CLIENTE ----
    add_styled_heading(doc, "2. Dados do Cliente", level=2)

    table2 = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    data2 = [
        ("Nome completo:", "{{ name }}"),
        ("Telefone:", "{{ phone }}"),
        ("E-mail:", "{{ email }}"),
        ("Cidade/Estado:", "{{ city_state }}"),
    ]
    for i, (label, value) in enumerate(data2):
        add_info_row(table2, label, value, i)

    # ---- RESUMO DO CASO ----
    add_styled_heading(doc, "3. Resumo do Caso", level=2)
    doc.add_paragraph("{{ case_summary }}")

    # ---- RESPOSTAS DA TRIAGEM ----
    add_styled_heading(doc, "4. Respostas da Triagem", level=2)

    doc.add_paragraph(
        "{% for answer in answers %}"
    )
    p = doc.add_paragraph()
    run = p.add_run("{{ answer.pergunta }}")
    run.bold = True
    doc.add_paragraph("{{ answer.resposta }}")
    doc.add_paragraph("")
    doc.add_paragraph("{% endfor %}")

    # ---- OBSERVAÇÕES DA IA ----
    add_styled_heading(doc, "5. Observações do Assistente Virtual", level=2)
    doc.add_paragraph("{{ ai_observations }}")

    # ---- RECOMENDAÇÃO ----
    add_styled_heading(doc, "6. Recomendação", level=2)
    doc.add_paragraph("{{ recommendation }}")

    # ---- RODAPÉ ----
    doc.add_paragraph("─" * 60)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "Documento gerado automaticamente pelo Assistente Virtual Breviliere\n"
        "Classificação: CONFIDENCIAL — Uso interno do escritório"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Salvar
    path = TEMPLATES_DIR / "briefing.docx"
    doc.save(str(path))
    print(f"  ✅ {path.name} criado")
    return path


# ============================================================
# TEMPLATE 2: PROPOSTA DE HONORÁRIOS
# ============================================================
def create_proposta_template():
    """Cria o template de proposta de honorários para o cliente."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # ---- CABEÇALHO ----
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("BREVILIERE ADVOCACIA")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Proposta de Prestação de Serviços Advocatícios")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_paragraph("")

    # ---- DESTINATÁRIO ----
    add_styled_heading(doc, "Prezado(a) {{ name }},", level=2)

    doc.add_paragraph(
        "Agradecemos a confiança em nosso escritório. Após análise preliminar "
        "do seu caso, apresentamos a seguinte proposta de honorários advocatícios "
        "para a prestação de serviços jurídicos na área de {{ area }}."
    )

    # ---- OBJETO ----
    add_styled_heading(doc, "1. Objeto", level=2)
    doc.add_paragraph(
        "Prestação de serviços advocatícios para {{ service_description }}, "
        "incluindo todas as diligências necessárias para a defesa dos interesses "
        "do(a) cliente, tanto na esfera administrativa quanto judicial."
    )

    # ---- SERVIÇOS INCLUSOS ----
    add_styled_heading(doc, "2. Serviços Inclusos", level=2)
    services = [
        "Análise detalhada do caso e documentação",
        "Elaboração de peças processuais (petição inicial, contestação, recursos)",
        "Acompanhamento de audiências e sessões de julgamento",
        "Representação junto a órgãos administrativos (se aplicável)",
        "Comunicação regular sobre andamento processual",
        "Atendimento via WhatsApp para dúvidas sobre o processo",
    ]
    for service in services:
        doc.add_paragraph(service, style="List Bullet")

    # ---- HONORÁRIOS ----
    add_styled_heading(doc, "3. Honorários", level=2)

    table = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fee_data = [
        ("Modalidade:", "{{ fee_type }}"),
        ("Valor/Percentual:", "{{ fee_value }}"),
        ("Entrada:", "{{ fee_upfront }}"),
        ("Condições:", "{{ fee_conditions }}"),
    ]
    for i, (label, value) in enumerate(fee_data):
        add_info_row(table, label, value, i)

    doc.add_paragraph("")

    p = doc.add_paragraph()
    run = p.add_run("Observação: ")
    run.bold = True
    p.add_run(
        "Os honorários acima não incluem custas processuais, taxas judiciais, "
        "despesas com perícia, certidões ou deslocamentos, que serão cobrados "
        "à parte mediante comprovação."
    )

    # ---- FORMA DE PAGAMENTO ----
    add_styled_heading(doc, "4. Forma de Pagamento", level=2)
    doc.add_paragraph("{{ payment_method }}")

    # ---- PRAZO DE VALIDADE ----
    add_styled_heading(doc, "5. Prazo de Validade", level=2)
    doc.add_paragraph(
        "Esta proposta tem validade de {{ validity_days }} dias a partir "
        "da data de emissão ({{ date }})."
    )

    # ---- ACEITE ----
    add_styled_heading(doc, "6. Aceite", level=2)
    doc.add_paragraph(
        "Caso esteja de acordo com os termos desta proposta, solicitamos "
        "a assinatura abaixo para formalização da contratação."
    )

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Linha de assinatura cliente
    sig1 = doc.add_paragraph()
    sig1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig1.add_run("_" * 50)
    name1 = doc.add_paragraph()
    name1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name1.add_run("{{ name }}")
    run.font.size = Pt(10)
    cpf1 = doc.add_paragraph()
    cpf1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cpf1.add_run("CPF: {{ cpf }}")
    run.font.size = Pt(9)

    doc.add_paragraph("")

    # Linha de assinatura advogado
    sig2 = doc.add_paragraph()
    sig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig2.add_run("_" * 50)
    name2 = doc.add_paragraph()
    name2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name2.add_run("{{ lawyer_name }}")
    run.font.size = Pt(10)
    oab2 = doc.add_paragraph()
    oab2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = oab2.add_run("OAB/{{ lawyer_state }} {{ lawyer_oab }}")
    run.font.size = Pt(9)

    # ---- RODAPÉ ----
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("{{ city_state }}, {{ date }}")
    run.font.size = Pt(10)

    path = TEMPLATES_DIR / "proposta_honorarios.docx"
    doc.save(str(path))
    print(f"  ✅ {path.name} criado")
    return path


# ============================================================
# TEMPLATE 3: CONTRATO DE PRESTAÇÃO DE SERVIÇOS
# ============================================================
def create_contrato_template():
    """Cria o template de contrato de prestação de serviços."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    # ---- TÍTULO ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("CONTRATO DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

    doc.add_paragraph("")

    # ---- PARTES ----
    add_styled_heading(doc, "DAS PARTES", level=2)

    doc.add_paragraph(
        "CONTRATANTE: {{ name }}, {{ nationality }}, {{ marital_status }}, "
        "{{ profession }}, portador(a) do CPF nº {{ cpf }} e RG nº {{ rg }}, "
        "residente e domiciliado(a) em {{ address }}, {{ city_state }}, "
        "telefone {{ phone }}, e-mail {{ email }}."
    )

    doc.add_paragraph(
        "CONTRATADO: {{ firm_name }}, sociedade de advogados inscrita na "
        "OAB/{{ firm_state }} sob nº {{ firm_oab }}, com sede em {{ firm_address }}, "
        "neste ato representada por {{ lawyer_name }}, advogado(a) inscrito(a) na "
        "OAB/{{ lawyer_state }} sob nº {{ lawyer_oab }}."
    )

    # ---- CLÁUSULA 1 ----
    add_styled_heading(doc, "CLÁUSULA 1ª — DO OBJETO", level=2)
    doc.add_paragraph(
        "O presente contrato tem por objeto a prestação de serviços advocatícios "
        "pelo CONTRATADO ao CONTRATANTE, consistentes em {{ service_description }}, "
        "na área de {{ area }}, incluindo todas as medidas judiciais e "
        "extrajudiciais necessárias para a defesa dos interesses do CONTRATANTE."
    )

    # ---- CLÁUSULA 2 ----
    add_styled_heading(doc, "CLÁUSULA 2ª — DOS HONORÁRIOS", level=2)
    doc.add_paragraph(
        "Pelos serviços prestados, o CONTRATANTE pagará ao CONTRATADO honorários "
        "advocatícios na modalidade {{ fee_type }}, no valor/percentual de "
        "{{ fee_value }}, sendo:"
    )
    doc.add_paragraph("a) Entrada: {{ fee_upfront }};", style="List Bullet")
    doc.add_paragraph("b) {{ fee_conditions }}.", style="List Bullet")

    doc.add_paragraph(
        "Parágrafo único: As custas processuais, taxas judiciais, emolumentos, "
        "despesas com perícia, certidões e deslocamentos NÃO estão incluídas nos "
        "honorários e serão de responsabilidade do CONTRATANTE, mediante "
        "comprovação prévia."
    )

    # ---- CLÁUSULA 3 ----
    add_styled_heading(doc, "CLÁUSULA 3ª — DAS OBRIGAÇÕES DO CONTRATADO", level=2)
    obligations_firm = [
        "Defender os interesses do CONTRATANTE com zelo, diligência e sigilo profissional;",
        "Manter o CONTRATANTE informado sobre o andamento do processo;",
        "Comparecer a audiências e praticar todos os atos processuais necessários;",
        "Disponibilizar canal de atendimento via WhatsApp para atualizações;",
        "Elaborar todas as peças processuais necessárias à defesa.",
    ]
    for ob in obligations_firm:
        doc.add_paragraph(ob, style="List Bullet")

    # ---- CLÁUSULA 4 ----
    add_styled_heading(doc, "CLÁUSULA 4ª — DAS OBRIGAÇÕES DO CONTRATANTE", level=2)
    obligations_client = [
        "Fornecer ao CONTRATADO todos os documentos e informações necessários;",
        "Comparecer a audiências quando convocado;",
        "Efetuar o pagamento dos honorários nas datas acordadas;",
        "Comunicar qualquer alteração de endereço, telefone ou e-mail;",
        "Não contratar outro profissional para o mesmo caso sem ciência do CONTRATADO.",
    ]
    for ob in obligations_client:
        doc.add_paragraph(ob, style="List Bullet")

    # ---- CLÁUSULA 5 ----
    add_styled_heading(doc, "CLÁUSULA 5ª — DA RESCISÃO", level=2)
    doc.add_paragraph(
        "O presente contrato poderá ser rescindido por qualquer das partes mediante "
        "notificação por escrito com antecedência mínima de 15 (quinze) dias. Em caso "
        "de rescisão, os honorários serão devidos proporcionalmente ao trabalho "
        "efetivamente realizado."
    )

    # ---- CLÁUSULA 6 ----
    add_styled_heading(doc, "CLÁUSULA 6ª — DA PROTEÇÃO DE DADOS (LGPD)", level=2)
    doc.add_paragraph(
        "O CONTRATADO se compromete a tratar os dados pessoais do CONTRATANTE "
        "exclusivamente para a finalidade da prestação dos serviços jurídicos "
        "contratados, em conformidade com a Lei Geral de Proteção de Dados "
        "(Lei nº 13.709/2018). Os dados serão armazenados pelo prazo de 5 (cinco) "
        "anos após o encerramento do processo, conforme obrigação legal."
    )

    # ---- CLÁUSULA 7 ----
    add_styled_heading(doc, "CLÁUSULA 7ª — DO FORO", level=2)
    doc.add_paragraph(
        "As partes elegem o foro da Comarca de {{ forum }}, Estado de {{ state }}, "
        "para dirimir quaisquer dúvidas ou litígios oriundos deste contrato, "
        "renunciando a qualquer outro, por mais privilegiado que seja."
    )

    # ---- ASSINATURAS ----
    doc.add_paragraph("")
    doc.add_paragraph(
        "E por estarem justos e contratados, firmam o presente instrumento em "
        "2 (duas) vias de igual teor e forma."
    )

    doc.add_paragraph("")
    date_line = doc.add_paragraph()
    date_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_line.add_run("{{ city_state }}, {{ date }}")

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Assinaturas lado a lado
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Contratante
    sig_table.rows[0].cells[0].text = "_" * 35
    sig_table.rows[1].cells[0].text = "{{ name }}"
    sig_table.rows[2].cells[0].text = "CPF: {{ cpf }}"
    for cell in [sig_table.rows[0].cells[0], sig_table.rows[1].cells[0], sig_table.rows[2].cells[0]]:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contratado
    sig_table.rows[0].cells[1].text = "_" * 35
    sig_table.rows[1].cells[1].text = "{{ lawyer_name }}"
    sig_table.rows[2].cells[1].text = "OAB/{{ lawyer_state }} {{ lawyer_oab }}"
    for cell in [sig_table.rows[0].cells[1], sig_table.rows[1].cells[1], sig_table.rows[2].cells[1]]:
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Testemunhas
    doc.add_paragraph("")
    add_styled_heading(doc, "Testemunhas:", level=3)
    wit_table = doc.add_table(rows=3, cols=2)
    wit_table.rows[0].cells[0].text = "_" * 35
    wit_table.rows[1].cells[0].text = "Nome:"
    wit_table.rows[2].cells[0].text = "CPF:"
    wit_table.rows[0].cells[1].text = "_" * 35
    wit_table.rows[1].cells[1].text = "Nome:"
    wit_table.rows[2].cells[1].text = "CPF:"

    path = TEMPLATES_DIR / "contrato_servicos.docx"
    doc.save(str(path))
    print(f"  ✅ {path.name} criado")
    return path


# ============================================================
# MAIN
# ============================================================
def main():
    os.makedirs(TEMPLATES_DIR, exist_ok=True)

    print("\n📄 Gerando templates DOCX profissionais...\n")

    create_briefing_template()
    create_proposta_template()
    create_contrato_template()

    print(f"\n✅ 3 templates criados em: {TEMPLATES_DIR}")
    print("   Placeholders usam sintaxe Jinja2: {{ variavel }}")
    print("   Compatível com docxtpl usado em generator.py\n")


if __name__ == "__main__":
    main()
