"""
Script para criar templates DOCX básicos para o DocumentGenerator.
Executar: python scripts/create_docx_templates.py
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def create_template_briefing():
    """Cria template de briefing jurídico."""
    doc = Document()
    
    # Título
    title = doc.add_heading('BRIEFING JURÍDICO - CLIENTE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Dados do cliente
    doc.add_heading('Dados do Cliente', level=1)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    dados = [
        ('Nome:', '{{ lead.name }}'),
        ('Telefone:', '{{ lead.phone }}'),
        ('Área Jurídica:', '{{ lead.area_juridica }}'),
        ('Score:', '{{ lead.score }}'),
        ('Data Atendimento:', '{{ data_atendimento }}')
    ]
    
    for i, (label, valor) in enumerate(dados):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = valor
    
    # Triagem
    doc.add_heading('Triagem Realizada', level=1)
    doc.add_paragraph('{{ triagem_resumo }}')
    
    # Respostas da triagem
    doc.add_heading('Detalhes da Triagem', level=2)
    for i in range(10):
        doc.add_paragraph(f'{{{{ pergunta_{i+1} }}}}')
        doc.add_paragraph(f'{{{{ resposta_{i+1} }}}}')
        doc.add_paragraph('')
    
    # Observações
    doc.add_heading('Observações do Advogado', level=1)
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    return doc

def create_template_proposta():
    """Cria template de proposta de honorários."""
    doc = Document()
    
    # Título
    title = doc.add_heading('PROPOSTA DE HONORÁRIOS ADVOCATÍCIOS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Dados
    doc.add_heading('Dados do Cliente', level=1)
    doc.add_paragraph('**Nome:** {{ lead.name }}')
    doc.add_paragraph('**Área:** {{ lead.area_juridica }}')
    
    # Objeto
    doc.add_heading('Objeto da Contratação', level=1)
    doc.add_paragraph('{{ descricao_servicos }}')
    
    # Honorários
    doc.add_heading('Honorários Advocatícios', level=1)
    doc.add_paragraph('**Valor:** {{ valor_honorarios }}')
    doc.add_paragraph('**Forma de Pagamento:** {{ forma_pagamento }}')
    
    # Condições
    doc.add_heading('Condições Gerais', level=1)
    doc.add_paragraph('{{ condicoes_gerais }}')
    
    # Assinaturas
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('São Paulo, {{ data_assinatura }}')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('___________________________')
    doc.add_paragraph('Advogado Responsável')
    doc.add_paragraph('OAB/SP nº XXXXX')
    
    return doc

def create_template_contrato():
    """Cria template de contrato de prestação de serviços."""
    doc = Document()
    
    # Título
    title = doc.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS ADVOCATÍCIOS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Partes
    doc.add_heading('Das Partes', level=1)
    doc.add_paragraph('**CONTRATANTE:** {{ cliente_nome }}, {{ cliente_nacionalidade }}, {{ cliente_estado_civil }}, {{ cliente_profissao }}, portador do RG nº {{ cliente_rg }}, inscrito no CPF sob nº {{ cliente_cpf }}, residente e domiciliado em {{ cliente_endereco }}.')
    doc.add_paragraph('**CONTRATADO:** {{ advogado_nome }}, advogado, inscrito na OAB/SP sob nº {{ advogado_oab }}, com escritório profissional em {{ advogado_endereco }}.')
    
    # Objeto
    doc.add_heading('Do Objeto', level=1)
    doc.add_paragraph('{{ objeto_contrato }}')
    
    # Obrigações
    doc.add_heading('Das Obrigações', level=1)
    doc.add_heading('Do Contratante', level=2)
    doc.add_paragraph('{{ obrigacoes_contratante }}')
    
    doc.add_heading('Do Contratado', level=2)
    doc.add_paragraph('{{ obrigacoes_contratado }}')
    
    # Honorários
    doc.add_heading('Dos Honorários', level=1)
    doc.add_paragraph('{{ honorarios_descricao }}')
    
    # Vigência
    doc.add_heading('Da Vigência', level=1)
    doc.add_paragraph('{{ vigencia }}')
    
    # Foro
    doc.add_heading('Do Foro', level=1)
    doc.add_paragraph('{{ foro }}')
    
    # Assinaturas
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('São Paulo, {{ data_assinatura }}')
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('___________________________')
    doc.add_paragraph('{{ cliente_nome }}')
    doc.add_paragraph('CONTRATANTE')
    doc.add_paragraph('')
    doc.add_paragraph('___________________________')
    doc.add_paragraph('{{ advogado_nome }}')
    doc.add_paragraph('CONTRATADO')
    
    return doc

if __name__ == "__main__":
    templates_dir = Path("src/documents/templates")
    templates_dir.mkdir(parents=True, exist_ok=True)
    
    # Criar templates
    templates = [
        ("briefing.docx", create_template_briefing),
        ("proposta.docx", create_template_proposta),
        ("contrato.docx", create_template_contrato)
    ]
    
    for filename, create_func in templates:
        doc = create_func()
        filepath = templates_dir / filename
        doc.save(str(filepath))
        print(f"✅ Template criado: {filepath}")
    
    print("\n📁 Templates criados com sucesso em src/documents/templates/")
